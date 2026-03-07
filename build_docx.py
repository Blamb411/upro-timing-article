"""
Build Word document for UPRO Timing Strategies article.
Generates individual charts from analysis data and assembles into .docx.

v2: Next-open execution, T-bill cash yield, synthetic pre-2009 UPRO,
    walk-forward testing, parameter robustness heatmap, expanded metrics.
"""

import os
import sys
import re
import numpy as np
import pandas as pd
import yfinance as yf
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.lines import Line2D
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Add analysis script directory to path
_analysis_dir = r"C:\Users\Admin\Trading\repos\spy-80-delta\Strategies\80-Delta Call Strategy"
sys.path.insert(0, _analysis_dir)

_this_dir = os.path.dirname(os.path.abspath(__file__))
_chart_dir = os.path.join(_this_dir, "charts")
os.makedirs(_chart_dir, exist_ok=True)

# ======================================================================
# PARAMETERS (match analysis script)
# ======================================================================
INITIAL_CAPITAL = 100_000
DATA_START = "2008-01-01"
END_DATE = "2026-03-03"
TRADING_DAYS_PER_YEAR = 252

COLORS = {
    "upro_bh": "#1f77b4",
    "vix": "#d62728",
    "dual": "#2ca02c",
    "hfea": "#ff7f0e",
    "dd_exit": "#9467bd",
    "composite": "#8c564b",
    "spy": "#555555",
    "syn3x": "#17becf",
    "syn3x_m": "#bcbd22",
    "static3x": "#e377c2",
    "sma_gate": "#e6550d",
}

# ======================================================================
# DATA LOADING
# ======================================================================

def load_data():
    """Download all needed market data. Returns DataFrames (Open+Close) for
    tradeable instruments, Series for VIX and IRX."""
    print("Downloading market data...")
    tickers = {
        "SPY": "SPY", "UPRO": "UPRO", "TMF": "TMF",
        "TLT": "TLT", "VIX": "^VIX", "IRX": "^IRX",
    }
    data = {}
    for name, ticker in tickers.items():
        df = yf.download(ticker, start=DATA_START, end=END_DATE, progress=False,
                         auto_adjust=True, multi_level_index=False)
        df = df.dropna(subset=["Close"])
        if name in ("VIX", "IRX"):
            data[name] = df["Close"].squeeze()
        else:
            data[name] = df[["Open", "Close"]].copy()
        print(f"  {name}: {len(data[name])} days")
    return data


def load_extended_spy():
    """Download SPY and IRX back to 1993 for synthetic UPRO testing."""
    print("  Loading extended SPY data (1993+)...")
    spy = yf.download("SPY", start="1993-01-29", end=END_DATE, progress=False,
                       auto_adjust=True, multi_level_index=False)
    spy = spy[["Open", "Close"]].dropna()
    print(f"    SPY extended: {len(spy)} days")

    irx = yf.download("^IRX", start="1993-01-29", end=END_DATE, progress=False,
                       auto_adjust=True, multi_level_index=False)
    irx = irx["Close"].squeeze().dropna()
    print(f"    IRX extended: {len(irx)} days")
    return spy, irx


# ======================================================================
# HELPERS
# ======================================================================

def get_daily_tbill_rate(irx_series):
    """Convert ^IRX (13-week T-bill yield, e.g., 5.2 = 5.2%) to daily return."""
    return (1 + irx_series / 100) ** (1 / 252) - 1


def avg_rf_annual(tbill_daily, dates):
    """Compute average annualized risk-free rate over a date range."""
    if tbill_daily is None or dates is None or len(dates) == 0:
        return 0.0
    tb = tbill_daily.reindex(dates).dropna()
    if len(tb) == 0:
        return 0.0
    return float(tb.mean()) * TRADING_DAYS_PER_YEAR


def compute_metrics(values, name, dates=None, trades=0, pct_invested=1.0,
                    rf_annual=0.0):
    """Compute expanded performance metrics.
    rf_annual: annualized risk-free rate (e.g. 0.02 for 2%) used for
               Sharpe, Sortino, and Calmar excess-return calculations.
    """
    vals = np.array(values, dtype=float)
    n_years = len(vals) / TRADING_DAYS_PER_YEAR
    cagr = (vals[-1] / vals[0]) ** (1.0 / n_years) - 1
    daily_rets = np.diff(vals) / vals[:-1]
    daily_rf = rf_annual / TRADING_DAYS_PER_YEAR
    excess_rets = daily_rets - daily_rf
    sharpe = (np.mean(excess_rets) / np.std(daily_rets) * np.sqrt(TRADING_DAYS_PER_YEAR)
              if np.std(daily_rets) > 0 else 0)
    neg_excess = excess_rets[excess_rets < 0]
    sortino = (np.mean(excess_rets) / np.std(neg_excess) * np.sqrt(TRADING_DAYS_PER_YEAR)
               if len(neg_excess) > 0 and np.std(neg_excess) > 0 else 0)
    cummax = np.maximum.accumulate(vals)
    drawdowns = vals / cummax - 1
    max_dd = drawdowns.min()

    # Calmar ratio (excess CAGR / |MaxDD|)
    excess_cagr = cagr - rf_annual
    calmar = excess_cagr / abs(max_dd) if max_dd != 0 else 0

    # Worst rolling 12-month return (252 trading days)
    worst_12m = 0.0
    if len(vals) > 252:
        rolling_12m = vals[252:] / vals[:-252] - 1
        worst_12m = float(rolling_12m.min())

    # Worst calendar month return
    worst_month = 0.0
    if dates is not None and len(dates) > 0:
        s = pd.Series(vals, index=dates)
        monthly = s.resample("ME").last()
        monthly_rets = monthly.pct_change().dropna()
        if len(monthly_rets) > 0:
            worst_month = float(monthly_rets.min())

    # Time underwater: longest streak (in days) below prior peak
    underwater_flags = drawdowns < 0
    time_underwater = 0
    current_streak = 0
    for uw in underwater_flags:
        if uw:
            current_streak += 1
            time_underwater = max(time_underwater, current_streak)
        else:
            current_streak = 0

    return {
        "name": name, "end_value": vals[-1], "cagr": cagr, "sharpe": sharpe,
        "sortino": sortino, "max_dd": max_dd, "calmar": calmar,
        "pct_invested": pct_invested,
        "num_trades": trades, "worst_12m": worst_12m, "worst_month": worst_month,
        "time_underwater": time_underwater,
    }


# ======================================================================
# STRATEGY RUNNERS — Next-Open Execution
#
# Pattern for all timing strategies:
#   Signal computed from close[i-1] (yesterday's close)
#   Trade executed at open[i] (today's open)
#   Portfolio marked at close[i] (today's close)
#   Cash periods earn T-bill daily rate
# ======================================================================

def run_upro_bh(upro_df, tbill_daily=None):
    """UPRO Buy & Hold: buy at first close, hold."""
    closes = upro_df["Close"].values
    portfolio_values = (INITIAL_CAPITAL / closes[0]) * closes
    dates = upro_df.index
    return dates, portfolio_values, compute_metrics(
        portfolio_values, "UPRO B&H", dates=dates, trades=1, pct_invested=1.0,
        rf_annual=avg_rf_annual(tbill_daily, dates))


def run_spy_bh(spy_close, start_date, tbill_daily=None):
    """SPY Buy & Hold (unchanged, close-only)."""
    spy = spy_close.loc[spy_close.index >= start_date]
    portfolio_values = (INITIAL_CAPITAL / spy.values[0]) * spy.values
    return spy.index, portfolio_values, compute_metrics(
        portfolio_values, "SPY B&H (1x)",
        rf_annual=avg_rf_annual(tbill_daily, spy.index))


def run_synthetic_3x(spy_close, start_date, rate=0.0, tbill_daily=None):
    """Synthetic daily-rebalanced 3x (unchanged, close-only)."""
    spy = spy_close.loc[spy_close.index >= start_date]
    prices = spy.values
    daily_rets = np.diff(prices) / prices[:-1]
    daily_borrow = (rate * 2.0) / TRADING_DAYS_PER_YEAR
    pv = np.zeros(len(prices))
    pv[0] = INITIAL_CAPITAL
    for i in range(1, len(prices)):
        pv[i] = pv[i - 1] * (1 + 3.0 * daily_rets[i - 1] - daily_borrow)
        if pv[i] <= 0:
            pv[i:] = 0
            break
    label = "no cost" if rate == 0 else f"{rate:.0%} margin"
    return spy.index, pv, compute_metrics(
        pv, f"Synthetic 3x ({label})",
        rf_annual=avg_rf_annual(tbill_daily, spy.index))


def run_static_3x(spy_close, start_date, rate=0.06, maint_margin=0.25, tbill_daily=None):
    """Static 3x: $100K equity + $200K borrowed, buy $300K SPY, hold."""
    spy = spy_close.loc[spy_close.index >= start_date]
    prices = spy.values
    dates = spy.index
    equity = INITIAL_CAPITAL
    debt = equity * 2.0
    shares = (equity + debt) / prices[0]
    daily_int = rate / TRADING_DAYS_PER_YEAR
    pv = np.zeros(len(prices))
    pv[0] = equity
    margin_called = False
    for i in range(1, len(prices)):
        if margin_called:
            pv[i] = pv[i - 1]; continue
        debt *= (1.0 + daily_int)
        pos_val = shares * prices[i]
        eq = pos_val - debt
        if eq <= 0 or eq < maint_margin * pos_val:
            pv[i] = max(0.0, pos_val - debt)
            margin_called = True
        else:
            pv[i] = eq
    return dates, pv, compute_metrics(
        pv, f"Static 3x ({rate:.0%} margin)",
        rf_annual=avg_rf_annual(tbill_daily, dates))


def run_vix_filter(upro_df, vix_series, threshold, tbill_daily):
    """VIX filter: invest in UPRO when VIX < threshold, else cash."""
    common = upro_df.index.intersection(vix_series.index)
    upro_open = upro_df.loc[common, "Open"].values
    upro_close = upro_df.loc[common, "Close"].values
    vix = vix_series.loc[common].values
    tbill = tbill_daily.reindex(common).fillna(0).values

    portfolio, shares, invested = INITIAL_CAPITAL, 0.0, False
    values = []
    trades, days_invested = 0, 0

    for i in range(len(upro_close)):
        if i == 0:
            values.append(portfolio)
            continue

        # Signal from yesterday's close
        want_in = vix[i - 1] < threshold

        # Execute at today's open
        if want_in and not invested:
            shares = portfolio / upro_open[i]
            invested = True
            trades += 1
        elif not want_in and invested:
            portfolio = shares * upro_open[i]
            shares = 0.0
            invested = False

        # Mark to close
        if invested:
            values.append(shares * upro_close[i])
            days_invested += 1
        else:
            portfolio *= (1 + tbill[i])
            values.append(portfolio)

    pct_inv = days_invested / max(len(values) - 1, 1)
    return common, np.array(values), compute_metrics(
        values, f"VIX<{threshold}", dates=common, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, common))


def run_dual_momentum(upro_df, spy_df, tlt_df, tbill_daily, lookback=252):
    """Dual momentum: UPRO when SPY momentum > 0 and > TLT, else cash."""
    common = upro_df.index.intersection(spy_df.index).intersection(tlt_df.index)
    upro_open = upro_df.loc[common, "Open"].values
    upro_close = upro_df.loc[common, "Close"].values
    spy_close = spy_df.loc[common, "Close"].values
    tlt_close = tlt_df.loc[common, "Close"].values
    tbill = tbill_daily.reindex(common).fillna(0).values

    portfolio, shares, invested = INITIAL_CAPITAL, 0.0, False
    values = []
    trades, days_invested = 0, 0

    for i in range(len(upro_close)):
        if i == 0:
            values.append(portfolio)
            continue

        # Need lookback+1 data points to compute signal from close[i-1]
        if i <= lookback:
            if invested:
                values.append(shares * upro_close[i])
                days_invested += 1
            else:
                portfolio *= (1 + tbill[i])
                values.append(portfolio)
            continue

        # Signal from yesterday's close
        spy_mom = spy_close[i - 1] / spy_close[i - 1 - lookback] - 1
        tlt_mom = tlt_close[i - 1] / tlt_close[i - 1 - lookback] - 1
        want_in = spy_mom > 0 and spy_mom > tlt_mom

        # Execute at today's open
        if want_in and not invested:
            shares = portfolio / upro_open[i]
            invested = True
            trades += 1
        elif not want_in and invested:
            portfolio = shares * upro_open[i]
            shares = 0.0
            invested = False

        if invested:
            values.append(shares * upro_close[i])
            days_invested += 1
        else:
            portfolio *= (1 + tbill[i])
            values.append(portfolio)

    pct_inv = days_invested / max(len(values) - 1, 1)
    return common, np.array(values), compute_metrics(
        values, "Dual Momentum", dates=common, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, common))


def run_hfea(upro_df, tmf_df, upro_wt=0.55, rebal_days=63, tbill_daily=None):
    """HFEA 55/45: always invested, rebalance at open every rebal_days."""
    common = upro_df.index.intersection(tmf_df.index)
    upro_open = upro_df.loc[common, "Open"].values
    upro_close = upro_df.loc[common, "Close"].values
    tmf_open = tmf_df.loc[common, "Open"].values
    tmf_close = tmf_df.loc[common, "Close"].values

    # Enter at first close
    upro_shares = (INITIAL_CAPITAL * upro_wt) / upro_close[0]
    tmf_shares = (INITIAL_CAPITAL * (1 - upro_wt)) / tmf_close[0]
    values = [INITIAL_CAPITAL]
    rebal_count = 0

    for i in range(1, len(upro_close)):
        # Rebalance at open if due
        if i % rebal_days == 0:
            port_val_open = upro_shares * upro_open[i] + tmf_shares * tmf_open[i]
            upro_shares = (port_val_open * upro_wt) / upro_open[i]
            tmf_shares = (port_val_open * (1 - upro_wt)) / tmf_open[i]
            rebal_count += 1

        # Mark to close
        port_val = upro_shares * upro_close[i] + tmf_shares * tmf_close[i]
        values.append(port_val)

    return common, np.array(values), compute_metrics(
        values, "HFEA 55/45", dates=common, trades=1 + rebal_count, pct_invested=1.0,
        rf_annual=avg_rf_annual(tbill_daily, common))


def run_dd_exit(upro_df, threshold, cool_days, tbill_daily):
    """Drawdown-triggered exit with cooling period. Next-open execution."""
    dates = upro_df.index
    upro_open = upro_df["Open"].values
    upro_close = upro_df["Close"].values
    tbill = tbill_daily.reindex(dates).fillna(0).values

    # Start invested at first close
    shares = INITIAL_CAPITAL / upro_close[0]
    portfolio = INITIAL_CAPITAL
    invested = True
    ath = upro_close[0]
    cool_counter = 0
    in_cool = False
    exit_signal = False
    enter_signal = False
    values = [INITIAL_CAPITAL]
    trades = 1
    days_invested = 0

    for i in range(1, len(upro_close)):
        # Execute pending signals at today's open
        if exit_signal:
            portfolio = shares * upro_open[i]
            shares = 0.0
            invested = False
            exit_signal = False
            in_cool = True
            cool_counter = 0
        elif enter_signal:
            shares = portfolio / upro_open[i]
            invested = True
            enter_signal = False
            ath = upro_open[i]  # reset ATH to entry price

        # Mark to close
        if invested:
            val = shares * upro_close[i]
            values.append(val)
            days_invested += 1

            # Update ATH and check for exit signal (execute tomorrow)
            ath = max(ath, upro_close[i])
            dd = upro_close[i] / ath - 1
            if dd < -threshold:
                exit_signal = True
                trades += 1
        else:
            # Cash: earn T-bill
            portfolio *= (1 + tbill[i])
            values.append(portfolio)

            # Check re-entry conditions (execute tomorrow)
            if in_cool:
                cool_counter += 1
                if cool_counter >= cool_days or upro_close[i] >= ath:
                    enter_signal = True
                    in_cool = False
                    trades += 1

    pct_inv = days_invested / max(len(values) - 1, 1)
    name = f"DD{int(threshold * 100)}%/Cool{cool_days}"
    return dates, np.array(values), compute_metrics(
        values, name, dates=dates, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, dates))


def run_dd_exit_sma_gate(upro_df, spy_df, threshold, cool_days, tbill_daily, sma_length=200):
    """DD exit with SMA re-entry gate.
    Exit: same as DD exit (threshold drawdown, exit at next open).
    Re-entry: cooling expired AND SPY > N-day SMA (both required), OR new ATH.
    """
    common = upro_df.index.intersection(spy_df.index)
    upro_open = upro_df.loc[common, "Open"].values
    upro_close = upro_df.loc[common, "Close"].values
    spy_close = spy_df.loc[common, "Close"].values
    tbill = tbill_daily.reindex(common).fillna(0).values
    sma200 = pd.Series(spy_close).rolling(sma_length).mean().values

    shares = INITIAL_CAPITAL / upro_close[0]
    portfolio = INITIAL_CAPITAL
    invested = True
    ath = upro_close[0]
    cool_counter = 0
    in_cool = False
    exit_signal = False
    enter_signal = False
    values = [INITIAL_CAPITAL]
    trades = 1
    days_invested = 0

    for i in range(1, len(upro_close)):
        # Execute pending signals at today's open
        if exit_signal:
            portfolio = shares * upro_open[i]
            shares = 0.0
            invested = False
            exit_signal = False
            in_cool = True
            cool_counter = 0
        elif enter_signal:
            shares = portfolio / upro_open[i]
            invested = True
            enter_signal = False
            ath = upro_open[i]

        if invested:
            val = shares * upro_close[i]
            values.append(val)
            days_invested += 1

            ath = max(ath, upro_close[i])
            dd = upro_close[i] / ath - 1
            if dd < -threshold:
                exit_signal = True
                trades += 1
        else:
            portfolio *= (1 + tbill[i])
            values.append(portfolio)

            if in_cool:
                cool_counter += 1
                new_ath = upro_close[i] >= ath
                cool_expired = cool_counter >= cool_days
                sma_ok = (i < len(sma200) and not np.isnan(sma200[i])
                          and spy_close[i] > sma200[i])

                if new_ath or (cool_expired and sma_ok):
                    enter_signal = True
                    in_cool = False
                    trades += 1

    pct_inv = days_invested / max(len(values) - 1, 1)
    name = f"DD{int(threshold * 100)}%/Cool{cool_days}+SMA{sma_length}"
    return common, np.array(values), compute_metrics(
        values, name, dates=common, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, common))


def run_sma_filter(upro_df, spy_df, tbill_daily, sma_length=200, exit_buffer=0.0):
    """Pure SMA-based strategy: SPY SMA drives entry/exit, UPRO is held.
    Entry: buy UPRO when SPY closes above SMA(N).
    Exit: sell UPRO when SPY closes below SMA(N) * (1 - exit_buffer).
    Cash earns T-bill rate. Next-open execution.
    """
    common = upro_df.index.intersection(spy_df.index)
    upro_open = upro_df.loc[common, "Open"].values
    upro_close = upro_df.loc[common, "Close"].values
    spy_close = spy_df.loc[common, "Close"].values
    tbill = tbill_daily.reindex(common).fillna(0).values
    sma = pd.Series(spy_close).rolling(sma_length).mean().values

    shares = 0.0
    portfolio = INITIAL_CAPITAL
    invested = False
    exit_signal = False
    enter_signal = False
    values = [INITIAL_CAPITAL]
    trades = 0
    days_invested = 0

    for i in range(1, len(upro_close)):
        # Execute pending signals at today's open
        if exit_signal:
            portfolio = shares * upro_open[i]
            shares = 0.0
            invested = False
            exit_signal = False
        elif enter_signal:
            shares = portfolio / upro_open[i]
            invested = True
            enter_signal = False

        if invested:
            val = shares * upro_close[i]
            values.append(val)
            days_invested += 1

            # Check exit: SPY below SMA * (1 - buffer)
            if (not np.isnan(sma[i])
                    and spy_close[i] < sma[i] * (1 - exit_buffer)):
                exit_signal = True
                trades += 1
        else:
            portfolio *= (1 + tbill[i])
            values.append(portfolio)

            # Check entry: SPY above SMA
            if (not np.isnan(sma[i])
                    and spy_close[i] > sma[i]):
                enter_signal = True
                trades += 1

    pct_inv = days_invested / max(len(values) - 1, 1)
    buf_str = f"/{int(exit_buffer*100)}%buf" if exit_buffer > 0 else ""
    name = f"SMA{sma_length}{buf_str}"
    return common, np.array(values), compute_metrics(
        values, name, dates=common, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, common))


def run_dd_exit_bond(upro_df, bond_df, threshold, cool_days, bond_name="TLT",
                     tbill_daily=None):
    """DD exit with a bond ETF as cash vehicle during cooling periods.
    Exit: sell UPRO at next open, buy bond ETF at same open.
    Re-entry: sell bond at next open, buy UPRO at same open.
    """
    common = upro_df.index.intersection(bond_df.index)
    upro_open = upro_df.loc[common, "Open"].values
    upro_close = upro_df.loc[common, "Close"].values
    bond_open = bond_df.loc[common, "Open"].values
    bond_close = bond_df.loc[common, "Close"].values

    upro_shares = INITIAL_CAPITAL / upro_close[0]
    bond_shares = 0.0
    portfolio = INITIAL_CAPITAL
    in_upro = True
    ath = upro_close[0]
    cool_counter = 0
    in_cool = False
    exit_signal = False
    enter_signal = False
    values = [INITIAL_CAPITAL]
    trades = 1
    days_in_upro = 0

    for i in range(1, len(upro_close)):
        if exit_signal:
            portfolio = upro_shares * upro_open[i]
            upro_shares = 0.0
            bond_shares = portfolio / bond_open[i]
            in_upro = False
            exit_signal = False
            in_cool = True
            cool_counter = 0
        elif enter_signal:
            portfolio = bond_shares * bond_open[i]
            bond_shares = 0.0
            upro_shares = portfolio / upro_open[i]
            in_upro = True
            enter_signal = False
            ath = upro_open[i]

        if in_upro:
            val = upro_shares * upro_close[i]
            values.append(val)
            days_in_upro += 1
            ath = max(ath, upro_close[i])
            dd = upro_close[i] / ath - 1
            if dd < -threshold:
                exit_signal = True
                trades += 1
        else:
            val = bond_shares * bond_close[i]
            values.append(val)
            if in_cool:
                cool_counter += 1
                if cool_counter >= cool_days or upro_close[i] >= ath:
                    enter_signal = True
                    in_cool = False
                    trades += 1

    pct_inv = days_in_upro / max(len(values) - 1, 1)
    name = f"DD{int(threshold*100)}%/Cool{cool_days}+{bond_name}"
    return common, np.array(values), compute_metrics(
        values, name, dates=common, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, common))


def run_composite(upro_df, spy_df, vix_series, min_signals, tbill_daily):
    """Composite: invest when >= min_signals of (SMA200, VIX<25, 63d momentum)."""
    common = upro_df.index.intersection(spy_df.index).intersection(vix_series.index)
    upro_open = upro_df.loc[common, "Open"].values
    upro_close = upro_df.loc[common, "Close"].values
    spy_close = spy_df.loc[common, "Close"].values
    vix = vix_series.loc[common].values
    tbill = tbill_daily.reindex(common).fillna(0).values
    sma200 = pd.Series(spy_close).rolling(200).mean().values

    portfolio, shares, invested = INITIAL_CAPITAL, 0.0, False
    values = []
    trades, days_invested = 0, 0

    for i in range(len(upro_close)):
        if i == 0:
            values.append(portfolio)
            continue

        # Need 200+ days for SMA
        if i <= 200:
            if invested:
                values.append(shares * upro_close[i])
                days_invested += 1
            else:
                portfolio *= (1 + tbill[i])
                values.append(portfolio)
            continue

        # Signal from yesterday's close
        cond_sma = spy_close[i - 1] > sma200[i - 1]
        cond_vix = vix[i - 1] < 25
        mom_start = max(0, i - 1 - 63)
        cond_mom = spy_close[i - 1] > spy_close[mom_start]
        n_signals = sum([cond_sma, cond_vix, cond_mom])
        want_in = n_signals >= min_signals

        # Execute at today's open
        if want_in and not invested:
            shares = portfolio / upro_open[i]
            invested = True
            trades += 1
        elif not want_in and invested:
            portfolio = shares * upro_open[i]
            shares = 0.0
            invested = False

        if invested:
            values.append(shares * upro_close[i])
            days_invested += 1
        else:
            portfolio *= (1 + tbill[i])
            values.append(portfolio)

    pct_inv = days_invested / max(len(values) - 1, 1)
    return common, np.array(values), compute_metrics(
        values, f"Composite {min_signals}of3", dates=common, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, common))


# ======================================================================
# EXIT SIGNAL DIAGNOSTIC
# ======================================================================

def exit_signal_diagnostic(upro_df, spy_df, vix_series, threshold=0.25,
                            cool_days=40):
    """Log SPY drawdown and VIX at each UPRO DD exit signal."""
    common = upro_df.index.intersection(spy_df.index).intersection(vix_series.index)
    upro_close = upro_df.loc[common, "Close"].values
    spy_close = spy_df.loc[common, "Close"].values
    vix = vix_series.loc[common].values
    dates = common

    ath_upro = upro_close[0]
    ath_spy = spy_close[0]
    invested = True
    in_cool = False
    cool_counter = 0
    exit_signal = False
    enter_signal = False
    exits = []

    for i in range(1, len(upro_close)):
        if exit_signal:
            invested = False
            exit_signal = False
            in_cool = True
            cool_counter = 0
        elif enter_signal:
            invested = True
            enter_signal = False
            ath_upro = upro_close[i]

        if invested:
            ath_upro = max(ath_upro, upro_close[i])
            ath_spy = max(ath_spy, spy_close[i])
            dd_upro = upro_close[i] / ath_upro - 1
            if dd_upro < -threshold:
                dd_spy = spy_close[i] / ath_spy - 1
                exits.append({
                    "date": dates[i],
                    "upro_close": upro_close[i],
                    "upro_ath": ath_upro,
                    "upro_dd": dd_upro,
                    "spy_close": spy_close[i],
                    "spy_ath": ath_spy,
                    "spy_dd": dd_spy,
                    "vix": vix[i],
                })
                exit_signal = True
        else:
            if in_cool:
                cool_counter += 1
                if cool_counter >= cool_days or upro_close[i] >= ath_upro:
                    enter_signal = True
                    in_cool = False

    return pd.DataFrame(exits)


# ======================================================================
# BLOCK BOOTSTRAP ROBUSTNESS TEST
# ======================================================================

def run_dd_exit_on_returns(daily_rets, threshold, cool_days, daily_rf=0.0):
    """Run DD exit strategy on a synthetic price path built from daily returns.
    Returns (values_array, bh_values_array)."""
    prices = np.zeros(len(daily_rets) + 1)
    prices[0] = 100.0
    for i, r in enumerate(daily_rets):
        prices[i + 1] = prices[i] * (1 + r)
        if prices[i + 1] <= 0:
            prices[i + 1:] = 0.001
            break

    # Buy-and-hold
    bh = (INITIAL_CAPITAL / prices[0]) * prices

    # DD exit strategy (simplified: close-only, no open/close distinction)
    shares = INITIAL_CAPITAL / prices[0]
    portfolio = INITIAL_CAPITAL
    invested = True
    ath = prices[0]
    cool_counter = 0
    in_cool = False
    exit_flag = False
    enter_flag = False
    values = [INITIAL_CAPITAL]

    for i in range(1, len(prices)):
        if exit_flag:
            portfolio = shares * prices[i]
            shares = 0.0
            invested = False
            exit_flag = False
            in_cool = True
            cool_counter = 0
        elif enter_flag:
            shares = portfolio / prices[i]
            invested = True
            enter_flag = False
            ath = prices[i]

        if invested:
            val = shares * prices[i]
            values.append(val)
            ath = max(ath, prices[i])
            dd = prices[i] / ath - 1
            if dd < -threshold:
                exit_flag = True
        else:
            portfolio *= (1 + daily_rf)
            values.append(portfolio)
            if in_cool:
                cool_counter += 1
                if cool_counter >= cool_days or prices[i] >= ath:
                    enter_flag = True
                    in_cool = False

    return np.array(values), bh


def run_block_bootstrap(upro_df, tbill_daily, threshold=0.25, cool_days=40,
                         block_size=20, n_sim=1000, seed=42):
    """Block bootstrap: shuffle 20-day blocks of UPRO returns, run strategy
    on each simulated path. Returns DataFrame of per-sim metrics."""
    closes = upro_df["Close"].values
    daily_rets = np.diff(closes) / closes[:-1]
    avg_rf = float(tbill_daily.reindex(upro_df.index).fillna(0).mean())
    n = len(daily_rets)

    # Build overlapping blocks
    blocks = [daily_rets[i:i + block_size] for i in range(n - block_size + 1)]

    rng = np.random.default_rng(seed)
    results = []

    for _ in range(n_sim):
        # Sample blocks with replacement to fill the original length
        path = []
        while len(path) < n:
            idx = rng.integers(0, len(blocks))
            path.extend(blocks[idx].tolist())
        path = path[:n]

        strat_vals, bh_vals = run_dd_exit_on_returns(
            path, threshold, cool_days, daily_rf=avg_rf)

        n_years = len(strat_vals) / TRADING_DAYS_PER_YEAR
        strat_cagr = (strat_vals[-1] / strat_vals[0]) ** (1.0 / n_years) - 1
        bh_cagr = (bh_vals[-1] / bh_vals[0]) ** (1.0 / n_years) - 1

        strat_rets = np.diff(strat_vals) / strat_vals[:-1]
        bh_rets = np.diff(bh_vals) / bh_vals[:-1]

        rf_annual = avg_rf * TRADING_DAYS_PER_YEAR
        daily_rf_rate = rf_annual / TRADING_DAYS_PER_YEAR
        strat_sharpe = ((np.mean(strat_rets - daily_rf_rate) / np.std(strat_rets)
                         * np.sqrt(TRADING_DAYS_PER_YEAR))
                        if np.std(strat_rets) > 0 else 0)
        bh_sharpe = ((np.mean(bh_rets - daily_rf_rate) / np.std(bh_rets)
                      * np.sqrt(TRADING_DAYS_PER_YEAR))
                     if np.std(bh_rets) > 0 else 0)

        strat_cm = np.maximum.accumulate(strat_vals)
        strat_dd = (strat_vals / strat_cm - 1).min()
        bh_cm = np.maximum.accumulate(bh_vals)
        bh_dd = (bh_vals / bh_cm - 1).min()

        results.append({
            "strat_sharpe": strat_sharpe, "bh_sharpe": bh_sharpe,
            "strat_cagr": strat_cagr, "bh_cagr": bh_cagr,
            "strat_dd": strat_dd, "bh_dd": bh_dd,
        })

    return pd.DataFrame(results)


# ======================================================================
# VOLATILITY-NORMALIZED DRAWDOWN RULE
# ======================================================================

def run_vol_normalized_dd(upro_df, tbill_daily, k=6, vol_window=20,
                           cool_days=40):
    """Exit when drawdown exceeds k * rolling vol. Next-open execution."""
    dates = upro_df.index
    upro_open = upro_df["Open"].values
    upro_close = upro_df["Close"].values
    tbill = tbill_daily.reindex(dates).fillna(0).values

    daily_rets = np.zeros(len(upro_close))
    daily_rets[1:] = np.diff(upro_close) / upro_close[:-1]
    rolling_vol = pd.Series(daily_rets).rolling(vol_window).std().values

    shares = INITIAL_CAPITAL / upro_close[0]
    portfolio = INITIAL_CAPITAL
    invested = True
    ath = upro_close[0]
    cool_counter = 0
    in_cool = False
    exit_signal = False
    enter_signal = False
    values = [INITIAL_CAPITAL]
    trades = 1
    days_invested = 0

    for i in range(1, len(upro_close)):
        if exit_signal:
            portfolio = shares * upro_open[i]
            shares = 0.0
            invested = False
            exit_signal = False
            in_cool = True
            cool_counter = 0
        elif enter_signal:
            shares = portfolio / upro_open[i]
            invested = True
            enter_signal = False
            ath = upro_open[i]

        if invested:
            val = shares * upro_close[i]
            values.append(val)
            days_invested += 1

            ath = max(ath, upro_close[i])
            dd = upro_close[i] / ath - 1

            # Dynamic threshold: k * rolling volatility
            if (not np.isnan(rolling_vol[i]) and rolling_vol[i] > 0
                    and dd < -(k * rolling_vol[i])):
                exit_signal = True
                trades += 1
        else:
            portfolio *= (1 + tbill[i])
            values.append(portfolio)
            if in_cool:
                cool_counter += 1
                if cool_counter >= cool_days or upro_close[i] >= ath:
                    enter_signal = True
                    in_cool = False
                    trades += 1

    pct_inv = days_invested / max(len(values) - 1, 1)
    name = f"VolNorm k={k}"
    return dates, np.array(values), compute_metrics(
        values, name, dates=dates, trades=trades, pct_invested=pct_inv,
        rf_annual=avg_rf_annual(tbill_daily, dates))


# ======================================================================
# SYNTHETIC PRE-2009 UPRO
# ======================================================================

def build_synthetic_upro(spy_df, expense_ratio=0.0091):
    """Build synthetic 3x daily-rebalanced ETF from extended SPY data.
    Derives Open from leveraged overnight gap for proper next-open execution."""
    spy_close = spy_df["Close"].values
    spy_open = spy_df["Open"].values

    # Daily close-to-close returns
    daily_rets = np.diff(spy_close) / spy_close[:-1]
    syn_rets = 3 * daily_rets - expense_ratio / 252

    # Build cumulative close series
    syn_close = np.zeros(len(spy_close))
    syn_close[0] = 100.0
    for i in range(1, len(spy_close)):
        syn_close[i] = syn_close[i - 1] * (1 + syn_rets[i - 1])
        if syn_close[i] <= 0:
            syn_close[i:] = 0.001
            break

    # Derive Open: syn_open[i] = syn_close[i-1] * (1 + 3 * overnight_spy_return)
    syn_open = np.zeros(len(spy_close))
    syn_open[0] = syn_close[0]
    for i in range(1, len(spy_close)):
        overnight_ret = spy_open[i] / spy_close[i - 1] - 1
        syn_open[i] = syn_close[i - 1] * (1 + 3 * overnight_ret)

    return pd.DataFrame({"Open": syn_open, "Close": syn_close}, index=spy_df.index)


# ======================================================================
# WALK-FORWARD TEST
# ======================================================================

def run_walk_forward(upro_df, vix_series, spy_df, tlt_df, tmf_df, tbill_daily):
    """Walk-forward: train on 2009-2016 (pick best DD params by Sharpe),
    test on 2017-2026.  Runs best params on the FULL dataset and splits
    the equity curve at the boundary so strategy state (position, ATH,
    cooling counter) carries over naturally."""
    print("\n  Walk-forward test...")

    train_end = "2016-12-31"
    test_start = "2017-01-01"

    upro_train = upro_df.loc[:train_end]
    tbill_train = tbill_daily.loc[:train_end]

    # In-sample grid search
    best_sharpe = -np.inf
    best_params = (0.25, 40)

    thresholds = [0.10, 0.15, 0.20, 0.25, 0.30]
    cool_days_list = [10, 20, 30, 40, 50, 60]

    for thresh in thresholds:
        for cool in cool_days_list:
            _, _, m = run_dd_exit(upro_train, thresh, cool, tbill_train)
            if m["sharpe"] > best_sharpe:
                best_sharpe = m["sharpe"]
                best_params = (thresh, cool)

    print(f"    In-sample best: DD{int(best_params[0] * 100)}%/Cool{best_params[1]} "
          f"(Sharpe={best_sharpe:.3f})")

    # Run best params on FULL dataset (state carries across boundary)
    full_dates, full_vals, _ = run_dd_exit(upro_df, best_params[0], best_params[1],
                                           tbill_daily)

    # Split at boundary to compute IS and OOS metrics separately
    split_idx = full_dates.searchsorted(pd.Timestamp(test_start))
    train_vals = full_vals[:split_idx]
    train_dates = full_dates[:split_idx]
    test_vals = full_vals[split_idx:]
    test_dates = full_dates[split_idx:]

    # Normalize test values to start at INITIAL_CAPITAL for fair CAGR comparison
    test_vals_norm = test_vals * (INITIAL_CAPITAL / test_vals[0])

    train_metrics = compute_metrics(
        train_vals, f"WF Train", dates=train_dates,
        rf_annual=avg_rf_annual(tbill_daily, train_dates))
    test_metrics = compute_metrics(
        test_vals_norm, f"WF Test", dates=test_dates,
        rf_annual=avg_rf_annual(tbill_daily, test_dates))

    print(f"    OOS: CAGR={test_metrics['cagr']:.1%}, "
          f"Sharpe={test_metrics['sharpe']:.3f}, MaxDD={test_metrics['max_dd']:.1%}")

    return best_params, train_metrics, test_metrics


# ======================================================================
# CHART GENERATION
# ======================================================================

def chart_equity_curves(bm_dates, bm_vals, best_strats, path):
    """Chart 1: Equity curves, log scale."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.semilogy(bm_dates, bm_vals, label="UPRO B&H", linewidth=2.5, color=COLORS["upro_bh"])
    strat_colors = ["vix", "dual", "hfea", "dd_exit", "composite"]
    for idx, (label, dates, vals, _) in enumerate(best_strats):
        c = COLORS[strat_colors[idx]] if idx < len(strat_colors) else "#333"
        ax.semilogy(dates, vals, label=label, linewidth=1.8, color=c, alpha=0.85)
    ax.set_title("Portfolio Value: UPRO B&H vs. Best Timing Strategies ($100K, Log Scale)",
                 fontweight="bold", fontsize=12)
    ax.set_ylabel("Portfolio Value ($)")
    ax.legend(loc="upper left", fontsize=9, framealpha=0.9)
    ax.grid(True, alpha=0.3, which="both")
    ax.xaxis.set_major_locator(mdates.YearLocator(2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.tick_params(axis="x", rotation=30)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_drawdowns(bm_dates, bm_vals, best_strats, path):
    """Chart 2: Drawdown comparison."""
    fig, ax = plt.subplots(figsize=(10, 5))
    bm_cm = np.maximum.accumulate(bm_vals)
    bm_dd = (bm_vals / bm_cm - 1) * 100
    ax.fill_between(bm_dates, bm_dd, 0, alpha=0.25, color=COLORS["upro_bh"], label="UPRO B&H")
    dd_label = "DD Exit"
    for label, dates, vals, _ in best_strats:
        if "DD Exit" in label or "DD" in label:
            v = np.array(vals)
            cm = np.maximum.accumulate(v)
            dd = (v / cm - 1) * 100
            ax.plot(dates, dd, label=label, linewidth=1.5, color=COLORS["dd_exit"])
            dd_label = label.split("(")[-1].rstrip(")") if "(" in label else label
    ax.set_title(f"Drawdown Over Time: UPRO B&H vs. {dd_label}", fontweight="bold", fontsize=12)
    ax.set_ylabel("Drawdown (%)")
    ax.legend(loc="lower left", fontsize=9)
    ax.grid(True, alpha=0.3)
    ax.xaxis.set_major_locator(mdates.YearLocator(2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.tick_params(axis="x", rotation=30)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_risk_return(bm_metrics, all_results, path):
    """Chart 3: Risk/return scatter."""
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.scatter(abs(bm_metrics["max_dd"]) * 100, bm_metrics["cagr"] * 100,
               s=200, color=COLORS["upro_bh"], marker="*", zorder=5, label="UPRO B&H")
    ax.annotate("UPRO B&H", (abs(bm_metrics["max_dd"]) * 100, bm_metrics["cagr"] * 100),
                textcoords="offset points", xytext=(8, 5), fontsize=8, fontweight="bold")

    def get_color(name):
        if "VIX<" in name: return COLORS["vix"]
        if "Dual" in name: return COLORS["dual"]
        if "HFEA" in name: return COLORS["hfea"]
        if "+SMA" in name: return COLORS["sma_gate"]
        if "DD" in name: return COLORS["dd_exit"]
        if "Composite" in name: return COLORS["composite"]
        return "#333"

    for m in all_results:
        c = get_color(m["name"])
        ax.scatter(abs(m["max_dd"]) * 100, m["cagr"] * 100, s=80, color=c, alpha=0.7, zorder=3)
        ax.annotate(m["name"], (abs(m["max_dd"]) * 100, m["cagr"] * 100),
                    textcoords="offset points", xytext=(5, 3), fontsize=7)

    legend_elements = [
        Line2D([0], [0], marker="*", color="w", markerfacecolor=COLORS["upro_bh"], markersize=12, label="UPRO B&H"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor=COLORS["vix"], markersize=8, label="VIX Filter"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor=COLORS["dual"], markersize=8, label="Dual Momentum"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor=COLORS["hfea"], markersize=8, label="HFEA"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor=COLORS["dd_exit"], markersize=8, label="DD Exit"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor=COLORS["sma_gate"], markersize=8, label="DD+SMA Gate"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor=COLORS["composite"], markersize=8, label="Composite"),
    ]
    ax.legend(handles=legend_elements, loc="lower left", fontsize=8)
    ax.set_xlabel("Max Drawdown (%)", fontsize=11)
    ax.set_ylabel("CAGR (%)", fontsize=11)
    ax.set_title("Risk vs. Return: All Strategy Variants", fontweight="bold", fontsize=12)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_leverage_comparison(spy_d, spy_v, s3_d, s3_v, s3m_d, s3m_v,
                              stat_d, stat_v, upro_d, upro_v, path):
    """Chart 4: Leverage comparison equity curves."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.semilogy(spy_d, spy_v, label="SPY B&H (1x)", linewidth=2, color=COLORS["spy"])
    ax.semilogy(s3_d, s3_v, label="Synthetic 3x (no cost)", linewidth=2, color=COLORS["syn3x"], linestyle="--")
    ax.semilogy(s3m_d, s3m_v, label="Synthetic 3x (6% margin)", linewidth=2, color=COLORS["syn3x_m"], linestyle="-.")
    ax.semilogy(stat_d, stat_v, label="Static 3x (6% margin)", linewidth=2, color=COLORS["static3x"], linestyle=":")
    ax.semilogy(upro_d, upro_v, label="UPRO B&H", linewidth=2.5, color=COLORS["upro_bh"])
    ax.set_title("Leverage Comparison: Four Ways to Get 3x Exposure ($100K, Log Scale)",
                 fontweight="bold", fontsize=12)
    ax.set_ylabel("Portfolio Value ($)")
    ax.legend(loc="upper left", fontsize=9, framealpha=0.9)
    ax.grid(True, alpha=0.3, which="both")
    ax.xaxis.set_major_locator(mdates.YearLocator(2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.tick_params(axis="x", rotation=30)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_dd_heatmap(upro_df, tbill_daily, path):
    """Chart 5: DD parameter robustness heatmap (Sharpe ratio)."""
    thresholds = [0.10, 0.15, 0.20, 0.25, 0.30]
    cool_days_list = [10, 20, 30, 40, 50, 60]

    sharpe_grid = np.zeros((len(thresholds), len(cool_days_list)))
    cagr_grid = np.zeros_like(sharpe_grid)
    dd_grid = np.zeros_like(sharpe_grid)

    for ti, thresh in enumerate(thresholds):
        for ci, cool in enumerate(cool_days_list):
            _, _, m = run_dd_exit(upro_df, thresh, cool, tbill_daily)
            sharpe_grid[ti, ci] = m["sharpe"]
            cagr_grid[ti, ci] = m["cagr"] * 100
            dd_grid[ti, ci] = m["max_dd"] * 100

    fig, ax = plt.subplots(figsize=(10, 7))
    im = ax.imshow(sharpe_grid, cmap="RdYlGn", aspect="auto")

    ax.set_xticks(range(len(cool_days_list)))
    ax.set_xticklabels(cool_days_list)
    ax.set_yticks(range(len(thresholds)))
    ax.set_yticklabels([f"{int(t * 100)}%" for t in thresholds])
    ax.set_xlabel("Cooling Period (days)")
    ax.set_ylabel("Drawdown Threshold")
    ax.set_title("DD Exit Strategy: Sharpe Ratio by Parameters\n(annotated: CAGR% / MaxDD%)",
                 fontweight="bold")

    for ti in range(len(thresholds)):
        for ci in range(len(cool_days_list)):
            s = sharpe_grid[ti, ci]
            c = cagr_grid[ti, ci]
            d = dd_grid[ti, ci]
            weight = "bold" if (thresholds[ti] == 0.25 and cool_days_list[ci] == 40) else "normal"
            ax.text(ci, ti, f"{s:.2f}\n{c:.0f}% / {d:.0f}%", ha="center", va="center",
                    fontsize=8, fontweight=weight)

    # Highlight DD25/Cool40 cell
    highlight_ti = thresholds.index(0.25)
    highlight_ci = cool_days_list.index(40)
    rect = plt.Rectangle((highlight_ci - 0.5, highlight_ti - 0.5), 1, 1,
                          linewidth=3, edgecolor="blue", facecolor="none")
    ax.add_patch(rect)

    plt.colorbar(im, label="Sharpe Ratio")
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_synthetic_upro(syn_bh_dates, syn_bh_vals, syn_dd_dates, syn_dd_vals, path):
    """Chart 6: Synthetic UPRO 2000-2009 — B&H vs DD25/Cool40."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.semilogy(syn_bh_dates, syn_bh_vals, label="Synthetic UPRO B&H",
                linewidth=2, color=COLORS["upro_bh"])
    ax.semilogy(syn_dd_dates, syn_dd_vals, label="DD25%/Cool40",
                linewidth=2, color=COLORS["dd_exit"])
    ax.set_title("Synthetic UPRO (2000-2009): B&H vs. DD25%/Cool40 ($100K, Log Scale)",
                 fontweight="bold", fontsize=12)
    ax.set_ylabel("Portfolio Value ($)")
    ax.legend(loc="upper left", fontsize=9)
    ax.grid(True, alpha=0.3, which="both")
    ax.xaxis.set_major_locator(mdates.YearLocator())
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.tick_params(axis="x", rotation=30)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_walk_forward(train_metrics, test_metrics, best_params, path):
    """Chart 7: Walk-forward in-sample vs OOS comparison."""
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    labels = ["CAGR", "Sharpe", "Max DD"]
    train_vals = [train_metrics["cagr"] * 100, train_metrics["sharpe"],
                  train_metrics["max_dd"] * 100]
    test_vals = [test_metrics["cagr"] * 100, test_metrics["sharpe"],
                 test_metrics["max_dd"] * 100]

    x = np.arange(len(labels))
    width = 0.35

    axes[0].bar(x - width / 2, train_vals, width, label="In-Sample (2009-2016)",
                color="#2ca02c", alpha=0.8)
    axes[0].bar(x + width / 2, test_vals, width, label="Out-of-Sample (2017-2026)",
                color="#d62728", alpha=0.8)
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(labels)
    axes[0].set_title(f"DD{int(best_params[0] * 100)}%/Cool{best_params[1]}: Walk-Forward",
                      fontweight="bold")
    axes[0].legend(fontsize=8)
    axes[0].grid(True, alpha=0.3)

    for i, (tv, tsv) in enumerate(zip(train_vals, test_vals)):
        offset = 1 if tv >= 0 else -3
        axes[0].text(i - width / 2, tv + offset, f"{tv:.1f}", ha="center", fontsize=8)
        offset = 1 if tsv >= 0 else -3
        axes[0].text(i + width / 2, tsv + offset, f"{tsv:.1f}", ha="center", fontsize=8)

    # Panel 2: text summary
    axes[1].axis("off")
    summary = (
        f"Walk-Forward Validation\n"
        f"{'=' * 35}\n\n"
        f"Best In-Sample Parameters:\n"
        f"  DD Threshold: {int(best_params[0] * 100)}%\n"
        f"  Cooling Period: {best_params[1]} days\n\n"
        f"In-Sample (2009-2016):\n"
        f"  CAGR: {train_metrics['cagr']:.1%}\n"
        f"  Sharpe: {train_metrics['sharpe']:.3f}\n"
        f"  Max DD: {train_metrics['max_dd']:.1%}\n\n"
        f"Out-of-Sample (2017-2026):\n"
        f"  CAGR: {test_metrics['cagr']:.1%}\n"
        f"  Sharpe: {test_metrics['sharpe']:.3f}\n"
        f"  Max DD: {test_metrics['max_dd']:.1%}"
    )
    axes[1].text(0.1, 0.9, summary, transform=axes[1].transAxes, fontsize=10,
                 verticalalignment="top", fontfamily="monospace")

    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_sma_gate(base_dates, base_vals, sma_results, path):
    """Chart 8: SMA re-entry gate comparison — DD25/Cool40 vs +SMA variants."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.semilogy(base_dates, base_vals, label="DD25%/Cool40 (no gate, Sharpe 0.90)",
                linewidth=2.5, color=COLORS["dd_exit"])
    sma_colors = {50: "#e6550d", 100: "#fd8d3c", 200: "#fdae6b"}
    sma_styles = {50: "-", 100: "--", 200: "-."}
    for sma_len in [50, 100, 200]:
        d, v, m = sma_results[sma_len]
        ax.semilogy(d, v, label=f"+SMA{sma_len} gate (Sharpe {m['sharpe']:.2f})",
                    linewidth=1.8, color=sma_colors[sma_len],
                    linestyle=sma_styles[sma_len])
    ax.set_title("SMA Re-Entry Gate: DD25%/Cool40 vs. SMA-Gated Variants ($100K, Log Scale)",
                 fontweight="bold", fontsize=12)
    ax.set_ylabel("Portfolio Value ($)")
    ax.legend(loc="upper left", fontsize=9, framealpha=0.9)
    ax.grid(True, alpha=0.3, which="both")
    ax.xaxis.set_major_locator(mdates.YearLocator(2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.tick_params(axis="x", rotation=30)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_sma_filter(bm_dates, bm_vals, dd_dates, dd_vals, sma_filter_results, path):
    """Chart 13: Pure SMA filter comparison — UPRO B&H vs DD25/Cool40 vs SMA filters."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.semilogy(bm_dates, bm_vals, label="UPRO B&H (Sharpe 0.80)",
                linewidth=2, color=COLORS["upro_bh"], alpha=0.5)
    ax.semilogy(dd_dates, dd_vals, label="DD25%/Cool40 (Sharpe 0.90)",
                linewidth=2.5, color=COLORS["dd_exit"])
    sma_colors = {
        (200, 0.0): "#e6550d", (100, 0.0): "#fd8d3c",
        (50, 0.0): "#fdae6b", (200, 0.02): "#756bb1",
    }
    sma_styles = {
        (200, 0.0): "-", (100, 0.0): "--",
        (50, 0.0): "-.", (200, 0.02): ":",
    }
    for key in [(200, 0.0), (100, 0.0), (50, 0.0), (200, 0.02)]:
        d, v, m = sma_filter_results[key]
        ax.semilogy(d, v, label=f"{m['name']} (Sharpe {m['sharpe']:.2f})",
                    linewidth=1.8, color=sma_colors[key],
                    linestyle=sma_styles[key])
    ax.set_title("Pure SMA Filter: SPY SMA Drives UPRO Entry/Exit ($100K, Log Scale)",
                 fontweight="bold", fontsize=12)
    ax.set_ylabel("Portfolio Value ($)")
    ax.legend(loc="upper left", fontsize=9, framealpha=0.9)
    ax.grid(True, alpha=0.3, which="both")
    ax.xaxis.set_major_locator(mdates.YearLocator(2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.tick_params(axis="x", rotation=30)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_bond_cash(bh_dates, bh_vals, tb_dates, tb_vals, bond_results, path):
    """Chart 9: Bond cash alternatives — T-bill vs TLT vs TMF during cooling."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.semilogy(bh_dates, bh_vals, label="UPRO B&H",
                linewidth=2, color=COLORS["upro_bh"], alpha=0.5)
    ax.semilogy(tb_dates, tb_vals, label="DD25/Cool40 (T-bill cash)",
                linewidth=2, color=COLORS["dd_exit"])
    bond_colors = {"TLT": "#2ca02c", "TMF": "#d62728"}
    bond_styles = {"TLT": "-", "TMF": "--"}
    for name in ["TLT", "TMF"]:
        d, v, m = bond_results[name]
        ax.semilogy(d, v, label=f"DD25/Cool40 ({name} cash, Sharpe {m['sharpe']:.2f})",
                    linewidth=2, color=bond_colors[name], linestyle=bond_styles[name])
    ax.set_title("DD25%/Cool40: Cash Vehicle Comparison ($100K, Log Scale)",
                 fontweight="bold", fontsize=12)
    ax.set_ylabel("Portfolio Value ($)")
    ax.legend(loc="upper left", fontsize=9, framealpha=0.9)
    ax.grid(True, alpha=0.3, which="both")
    ax.xaxis.set_major_locator(mdates.YearLocator(2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.tick_params(axis="x", rotation=30)
    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


def chart_bootstrap(bootstrap_df, path):
    """Chart 14: Block bootstrap Sharpe distribution."""
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    # Panel 1: Sharpe distribution
    axes[0].hist(bootstrap_df["bh_sharpe"], bins=40, alpha=0.5, color=COLORS["upro_bh"],
                 label="UPRO B&H", density=True)
    axes[0].hist(bootstrap_df["strat_sharpe"], bins=40, alpha=0.5, color=COLORS["dd_exit"],
                 label="DD25%/Cool40", density=True)
    axes[0].axvline(bootstrap_df["bh_sharpe"].median(), color=COLORS["upro_bh"],
                    linestyle="--", linewidth=2)
    axes[0].axvline(bootstrap_df["strat_sharpe"].median(), color=COLORS["dd_exit"],
                    linestyle="--", linewidth=2)
    axes[0].set_xlabel("Sharpe Ratio")
    axes[0].set_ylabel("Density")
    axes[0].set_title("Block Bootstrap: Sharpe Distribution (1,000 sims)", fontweight="bold")
    axes[0].legend(fontsize=9)
    axes[0].grid(True, alpha=0.3)

    # Panel 2: Max drawdown distribution
    axes[1].hist(bootstrap_df["bh_dd"] * 100, bins=40, alpha=0.5, color=COLORS["upro_bh"],
                 label="UPRO B&H", density=True)
    axes[1].hist(bootstrap_df["strat_dd"] * 100, bins=40, alpha=0.5, color=COLORS["dd_exit"],
                 label="DD25%/Cool40", density=True)
    axes[1].axvline(bootstrap_df["bh_dd"].median() * 100, color=COLORS["upro_bh"],
                    linestyle="--", linewidth=2)
    axes[1].axvline(bootstrap_df["strat_dd"].median() * 100, color=COLORS["dd_exit"],
                    linestyle="--", linewidth=2)
    axes[1].set_xlabel("Max Drawdown (%)")
    axes[1].set_ylabel("Density")
    axes[1].set_title("Block Bootstrap: Max Drawdown Distribution", fontweight="bold")
    axes[1].legend(fontsize=9)
    axes[1].grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"  Saved: {path}")


# ======================================================================
# WORD DOCUMENT BUILDER
# ======================================================================

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_table_from_rows(doc, headers, rows, highlight_rows=None, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    if highlight_rows and i in highlight_rows:
                        run.bold = True

    return table


def build_document(article_md, chart_paths):
    """Parse markdown and build Word doc with charts."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Style modifications
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # ---- Parse markdown and build doc ----
    lines = article_md.split("\n")
    i = 0
    in_table = False
    table_rows = []
    table_headers = []

    # Chart placement: now handled by ![...]() image references in markdown
    chart_after_section = {}

    last_heading = None

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if line.strip() == "---":
            i += 1
            continue

        # Table handling
        if "|" in line and not line.strip().startswith("|--"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                # Check if next line is separator
                if i + 1 < len(lines) and re.match(r"^\|[-\s|]+\|$", lines[i + 1].strip()):
                    in_table = True
                    table_headers = cells
                    i += 2  # skip header + separator
                    table_rows = []
                    continue
                else:
                    in_table = True
                    table_headers = cells
                    i += 1
                    table_rows = []
                    continue
            else:
                table_rows.append(cells)
                i += 1
                continue
        elif in_table:
            # End of table - flush it
            in_table = False
            highlights = set()
            for ri, row in enumerate(table_rows):
                if any("**" in cell for cell in row):
                    highlights.add(ri)
            clean_rows = []
            for row in table_rows:
                clean_rows.append([c.replace("**", "") for c in row])
            clean_headers = [h.replace("**", "") for h in table_headers]

            add_table_from_rows(doc, clean_headers, clean_rows,
                                highlight_rows=highlights if highlights else None)
            doc.add_paragraph()

            if last_heading in chart_after_section:
                chart_key = chart_after_section[last_heading]
                if isinstance(chart_key, str) and chart_key in chart_paths:
                    if os.path.exists(chart_paths[chart_key]):
                        doc.add_picture(chart_paths[chart_key], width=Inches(6))
                        last_p = doc.paragraphs[-1]
                        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                elif isinstance(chart_key, list):
                    for ck in chart_key:
                        if ck in chart_paths and os.path.exists(chart_paths[ck]):
                            doc.add_picture(chart_paths[ck], width=Inches(6))
                            last_p = doc.paragraphs[-1]
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            doc.add_paragraph()
                del chart_after_section[last_heading]

            table_rows = []
            table_headers = []
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("##"):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("## "):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=1)
            last_heading = heading_text
            i += 1
            continue

        if line.startswith("### "):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=2)
            last_heading = heading_text
            i += 1
            if heading_text in chart_after_section:
                pass
            continue

        # Empty lines
        if line.strip() == "":
            i += 1
            continue

        # Bullet points
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r"^(\d+)\.\s+", line.strip())
        if num_match:
            text = line.strip()[num_match.end():]
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
            i += 1
            continue

        # Italics-only lines (disclosure, etc.)
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            text = line.strip().strip("*")
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Image references: ![alt](path)
        img_match = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
        if img_match:
            img_rel_path = img_match.group(1)
            img_path = os.path.join(_this_dir, img_rel_path)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(6))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line.strip())
        i += 1

    # Flush any remaining table
    if in_table and table_rows:
        clean_rows = [[c.replace("**", "") for c in row] for row in table_rows]
        clean_headers = [h.replace("**", "") for h in table_headers]
        highlights = {ri for ri, row in enumerate(table_rows) if any("**" in c for c in row)}
        add_table_from_rows(doc, clean_headers, clean_rows,
                            highlight_rows=highlights if highlights else None)

    return doc


def _add_formatted_text(paragraph, text):
    """Parse markdown bold/italic and add formatted runs to a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            sub_parts = re.split(r'(\*[^*]+?\*)', part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    if sp:
                        paragraph.add_run(sp)


# ======================================================================
# MAIN
# ======================================================================

def main():
    print("=" * 60)
    print("  Building UPRO Article Word Document (v2: next-open)")
    print("=" * 60)

    # ------------------------------------------------------------------
    # 1. Load data
    # ------------------------------------------------------------------
    data = load_data()
    spy_df = data["SPY"]        # DataFrame[Open, Close]
    upro_df = data["UPRO"]      # DataFrame[Open, Close]
    tmf_df = data["TMF"]        # DataFrame[Open, Close]
    tlt_df = data["TLT"]        # DataFrame[Open, Close]
    vix_series = data["VIX"]    # Series
    irx_series = data["IRX"]    # Series

    # ------------------------------------------------------------------
    # 2. Compute T-bill daily rates
    # ------------------------------------------------------------------
    tbill_daily = get_daily_tbill_rate(irx_series)

    # ------------------------------------------------------------------
    # 3. Close series for leverage comparison (backward compat)
    # ------------------------------------------------------------------
    spy_close = spy_df["Close"]

    # ------------------------------------------------------------------
    # 4. Run all strategies
    # ------------------------------------------------------------------
    print("\nRunning strategies...")

    # Benchmark
    bm_d, bm_v, bm_m = run_upro_bh(upro_df, tbill_daily)

    # Leverage comparison (unchanged, close-only)
    spy_d, spy_v, spy_m = run_spy_bh(spy_close, upro_df.index[0], tbill_daily)
    s3_d, s3_v, s3_m = run_synthetic_3x(spy_close, upro_df.index[0], 0.0, tbill_daily)
    s3m_d, s3m_v, s3m_m = run_synthetic_3x(spy_close, upro_df.index[0], 0.06, tbill_daily)
    stat_d, stat_v, stat_m = run_static_3x(spy_close, upro_df.index[0], 0.06,
                                            tbill_daily=tbill_daily)

    # Timing strategies
    all_results = []
    best_per_strategy = []

    # VIX filter
    print("  VIX filter...")
    vix_results = {}
    for t in [15, 20, 25, 30]:
        d, v, m = run_vix_filter(upro_df, vix_series, t, tbill_daily)
        vix_results[t] = (d, v, m)
        all_results.append(m)
    best_vix = max(vix_results.values(), key=lambda x: x[2]["sharpe"])
    best_per_strategy.append((f"VIX ({best_vix[2]['name']})", best_vix[0], best_vix[1], best_vix[2]))

    # Dual Momentum
    print("  Dual Momentum...")
    dm_d, dm_v, dm_m = run_dual_momentum(upro_df, spy_df, tlt_df, tbill_daily)
    all_results.append(dm_m)
    best_per_strategy.append(("Dual Momentum", dm_d, dm_v, dm_m))

    # HFEA
    print("  HFEA...")
    hf_d, hf_v, hf_m = run_hfea(upro_df, tmf_df, tbill_daily=tbill_daily)
    all_results.append(hf_m)
    best_per_strategy.append(("HFEA 55/45", hf_d, hf_v, hf_m))

    # Drawdown Exit (original 4x3 grid)
    print("  DD Exit grid...")
    dd_results = {}
    for thresh in [0.10, 0.15, 0.20, 0.25]:
        for cool in [20, 40, 60]:
            d, v, m = run_dd_exit(upro_df, thresh, cool, tbill_daily)
            dd_results[(thresh, cool)] = (d, v, m)
            all_results.append(m)
    best_dd = max(dd_results.values(), key=lambda x: x[2]["sharpe"])
    best_per_strategy.append((f"DD Exit ({best_dd[2]['name']})", best_dd[0], best_dd[1], best_dd[2]))

    # Composite
    print("  Composite...")
    comp_results = {}
    for n in [2, 3]:
        d, v, m = run_composite(upro_df, spy_df, vix_series, n, tbill_daily)
        comp_results[n] = (d, v, m)
        all_results.append(m)
    best_comp = max(comp_results.values(), key=lambda x: x[2]["sharpe"])
    best_per_strategy.append((f"Composite ({best_comp[2]['name']})", best_comp[0], best_comp[1], best_comp[2]))

    print("  All core strategies complete.")

    # ------------------------------------------------------------------
    # 5. DD25/Cool40 + SMA gate variants (50, 100, 200-day)
    # ------------------------------------------------------------------
    print("  DD25/Cool40 + SMA gate variants...")
    sma_gate_results = {}
    for sma_len in [50, 100, 200]:
        d, v, m = run_dd_exit_sma_gate(upro_df, spy_df, 0.25, 40, tbill_daily, sma_length=sma_len)
        sma_gate_results[sma_len] = (d, v, m)
        all_results.append(m)
        print(f"    SMA{sma_len} gate: CAGR={m['cagr']:.1%}, Sharpe={m['sharpe']:.3f}, "
              f"MaxDD={m['max_dd']:.1%}, %Inv={m['pct_invested']:.0%}")

    # ------------------------------------------------------------------
    # 5b. Pure SMA filter strategies (SPY SMA drives UPRO entry/exit)
    # ------------------------------------------------------------------
    print("  Pure SMA filter strategies...")
    sma_filter_results = {}
    sma_filter_variants = [
        (200, 0.0),    # SMA200 simple crossover
        (100, 0.0),    # SMA100 simple crossover
        (50,  0.0),    # SMA50 simple crossover
        (200, 0.02),   # SMA200 with 2% buffer (80-delta style)
    ]
    for sma_len, buf in sma_filter_variants:
        d, v, m = run_sma_filter(upro_df, spy_df, tbill_daily,
                                 sma_length=sma_len, exit_buffer=buf)
        sma_filter_results[(sma_len, buf)] = (d, v, m)
        all_results.append(m)
        print(f"    {m['name']}: CAGR={m['cagr']:.1%}, Sharpe={m['sharpe']:.3f}, "
              f"MaxDD={m['max_dd']:.1%}, %Inv={m['pct_invested']:.0%}, Trades={m['num_trades']}")

    # ------------------------------------------------------------------
    # 5c. DD25/Cool40 + bond cash alternatives (TLT, TMF)
    # ------------------------------------------------------------------
    print("  DD25/Cool40 + bond cash alternatives...")
    bond_cash_results = {}
    for bond_name, bond_df in [("TLT", tlt_df), ("TMF", tmf_df)]:
        d, v, m = run_dd_exit_bond(upro_df, bond_df, 0.25, 40, bond_name,
                                    tbill_daily=tbill_daily)
        bond_cash_results[bond_name] = (d, v, m)
        all_results.append(m)
        print(f"    {bond_name} cash: CAGR={m['cagr']:.1%}, Sharpe={m['sharpe']:.3f}, "
              f"MaxDD={m['max_dd']:.1%}")

    # ------------------------------------------------------------------
    # 6. Synthetic pre-2009 UPRO
    # ------------------------------------------------------------------
    print("\n  Synthetic UPRO...")
    ext_spy, ext_irx = load_extended_spy()
    ext_tbill = get_daily_tbill_rate(ext_irx)
    syn_upro = build_synthetic_upro(ext_spy)

    # 2000-2009 subset
    syn_2000 = syn_upro.loc["2000-01-01":"2009-06-25"]
    ext_tbill_2000 = ext_tbill.loc["2000-01-01":"2009-06-25"]
    syn_bh_d, syn_bh_v, syn_bh_m = run_upro_bh(syn_2000, ext_tbill_2000)
    syn_dd_d, syn_dd_v, syn_dd_m = run_dd_exit(syn_2000, 0.25, 40, ext_tbill_2000)
    print(f"    Synthetic 2000-2009 B&H: CAGR={syn_bh_m['cagr']:.1%}, MaxDD={syn_bh_m['max_dd']:.1%}")
    print(f"    Synthetic 2000-2009 DD25/Cool40: CAGR={syn_dd_m['cagr']:.1%}, MaxDD={syn_dd_m['max_dd']:.1%}")

    # Full 1993-2026
    syn_full_bh_d, syn_full_bh_v, syn_full_bh_m = run_upro_bh(syn_upro, ext_tbill)
    syn_full_dd_d, syn_full_dd_v, syn_full_dd_m = run_dd_exit(syn_upro, 0.25, 40, ext_tbill)
    print(f"    Synthetic full B&H: CAGR={syn_full_bh_m['cagr']:.1%}")
    print(f"    Synthetic full DD25/Cool40: CAGR={syn_full_dd_m['cagr']:.1%}, "
          f"Sharpe={syn_full_dd_m['sharpe']:.3f}")

    # ------------------------------------------------------------------
    # 7. Walk-forward test
    # ------------------------------------------------------------------
    wf_params, wf_train, wf_test = run_walk_forward(
        upro_df, vix_series, spy_df, tlt_df, tmf_df, tbill_daily)

    # ------------------------------------------------------------------
    # 7b. Exit signal diagnostic
    # ------------------------------------------------------------------
    print("\n  Exit signal diagnostic...")
    exit_diag = exit_signal_diagnostic(upro_df, spy_df, vix_series)
    diag_csv = os.path.join(_this_dir, "exit_diagnostics.csv")
    exit_diag.to_csv(diag_csv, index=False, float_format="%.4f")
    print(f"    {len(exit_diag)} exits logged")
    print(f"    Median SPY DD at exit: {exit_diag['spy_dd'].median():.1%}")
    print(f"    Median VIX at exit: {exit_diag['vix'].median():.1f}")
    print(f"    SPY DD range: {exit_diag['spy_dd'].min():.1%} to {exit_diag['spy_dd'].max():.1%}")
    print(f"    VIX range: {exit_diag['vix'].min():.1f} to {exit_diag['vix'].max():.1f}")

    # ------------------------------------------------------------------
    # 7c. Block bootstrap robustness test
    # ------------------------------------------------------------------
    print("\n  Block bootstrap (1,000 sims, 20-day blocks)...")
    bootstrap_df = run_block_bootstrap(upro_df, tbill_daily)
    pct_strat_wins = (bootstrap_df["strat_sharpe"] > bootstrap_df["bh_sharpe"]).mean()
    median_sharpe_imp = (bootstrap_df["strat_sharpe"] - bootstrap_df["bh_sharpe"]).median()
    median_dd_imp = (bootstrap_df["bh_dd"] - bootstrap_df["strat_dd"]).median()
    print(f"    Strategy Sharpe > B&H in {pct_strat_wins:.0%} of simulations")
    print(f"    Median Sharpe improvement: {median_sharpe_imp:+.3f}")
    print(f"    Median max DD reduction: {median_dd_imp:.1%}")

    # ------------------------------------------------------------------
    # 7d. Volatility-normalized drawdown rule
    # ------------------------------------------------------------------
    print("\n  Volatility-normalized DD rule...")
    vol_norm_results = {}
    for k_val in [4, 5, 6, 7]:
        d, v, m = run_vol_normalized_dd(upro_df, tbill_daily, k=k_val)
        vol_norm_results[k_val] = (d, v, m)
        print(f"    k={k_val}: CAGR={m['cagr']:.1%}, Sharpe={m['sharpe']:.3f}, "
              f"MaxDD={m['max_dd']:.1%}, Trades={m['num_trades']}")

    # ------------------------------------------------------------------
    # 8. Generate charts
    # ------------------------------------------------------------------
    print("\nGenerating charts...")
    chart_paths = {
        "equity_bh": os.path.join(_chart_dir, "01_equity_curves.png"),
        "leverage": os.path.join(_chart_dir, "02_leverage_comparison.png"),
        "drawdowns": os.path.join(_chart_dir, "03_drawdowns.png"),
        "risk_return": os.path.join(_chart_dir, "04_risk_return.png"),
        "dd_heatmap": os.path.join(_chart_dir, "05_dd_heatmap.png"),
        "synthetic_upro": os.path.join(_chart_dir, "06_synthetic_upro.png"),
        "walk_forward": os.path.join(_chart_dir, "07_walk_forward.png"),
        "sma_gate": os.path.join(_chart_dir, "08_sma_gate.png"),
        "bond_cash": os.path.join(_chart_dir, "09_bond_cash.png"),
        "sma_filter": os.path.join(_chart_dir, "13_sma_filter.png"),
        "bootstrap": os.path.join(_chart_dir, "14_bootstrap.png"),
    }
    # Keep old chart paths for document builder if files exist on disk
    for key, fname in [("margin_calls", "leverage_margin_calls_summary.png"),
                       ("leverage_full", "leverage_full.png")]:
        fpath = os.path.join(_chart_dir, fname)
        if os.path.exists(fpath):
            chart_paths[key] = fpath

    chart_equity_curves(bm_d, bm_v, best_per_strategy, chart_paths["equity_bh"])
    chart_leverage_comparison(spy_d, spy_v, s3_d, s3_v, s3m_d, s3m_v,
                              stat_d, stat_v, bm_d, bm_v, chart_paths["leverage"])
    chart_drawdowns(bm_d, bm_v, best_per_strategy, chart_paths["drawdowns"])
    chart_risk_return(bm_m, all_results, chart_paths["risk_return"])
    print("  DD heatmap (expanded grid)...")
    chart_dd_heatmap(upro_df, tbill_daily, chart_paths["dd_heatmap"])
    chart_synthetic_upro(syn_bh_d, syn_bh_v, syn_dd_d, syn_dd_v, chart_paths["synthetic_upro"])
    chart_walk_forward(wf_train, wf_test, wf_params, chart_paths["walk_forward"])
    dd25_40_data = dd_results[(0.25, 40)]
    chart_sma_gate(dd25_40_data[0], dd25_40_data[1], sma_gate_results,
                   chart_paths["sma_gate"])
    chart_bond_cash(bm_d, bm_v, dd25_40_data[0], dd25_40_data[1],
                    bond_cash_results, chart_paths["bond_cash"])
    chart_sma_filter(bm_d, bm_v, dd25_40_data[0], dd25_40_data[1],
                     sma_filter_results, chart_paths["sma_filter"])
    chart_bootstrap(bootstrap_df, chart_paths["bootstrap"])

    # ------------------------------------------------------------------
    # 9. Export CSV with expanded metrics
    # ------------------------------------------------------------------
    csv_path = os.path.join(_this_dir, "strategy_results.csv")
    all_metrics = [bm_m, spy_m, s3_m, s3m_m, stat_m] + all_results
    # Add synthetic and walk-forward results
    all_metrics.extend([
        {**syn_bh_m, "name": "Syn UPRO B&H (2000-09)"},
        {**syn_dd_m, "name": "Syn DD25/Cool40 (2000-09)"},
        {**syn_full_bh_m, "name": "Syn UPRO B&H (1993-26)"},
        {**syn_full_dd_m, "name": "Syn DD25/Cool40 (1993-26)"},
        {**wf_train, "name": f"WF Train DD{int(wf_params[0]*100)}/Cool{wf_params[1]}"},
        {**wf_test, "name": f"WF Test DD{int(wf_params[0]*100)}/Cool{wf_params[1]}"},
    ])
    for bname, (_, _, bm) in bond_cash_results.items():
        all_metrics.append({**bm, "name": f"DD25/Cool40+{bname} cash"})
    df_results = pd.DataFrame(all_metrics)
    df_results.to_csv(csv_path, index=False, float_format="%.6f")
    print(f"\nCSV saved: {csv_path}")

    # Print summary table
    print("\n" + "=" * 80)
    print(f"  {'Strategy':<30s} {'End Value':>12s} {'CAGR':>8s} {'Sharpe':>8s} "
          f"{'MaxDD':>8s} {'Calmar':>8s} {'%Inv':>6s}")
    print("-" * 80)
    for m in [bm_m] + all_results:
        print(f"  {m['name']:<30s} ${m['end_value']:>11,.0f} {m['cagr']:>7.1%} "
              f"{m['sharpe']:>7.3f} {m['max_dd']:>7.1%} {m['calmar']:>7.2f} "
              f"{m['pct_invested']:>5.0%}")
    print("=" * 80)

    # ------------------------------------------------------------------
    # 10. Build Word document
    # ------------------------------------------------------------------
    article_path = os.path.join(_this_dir, "article_draft.md")
    if os.path.exists(article_path):
        print("\nBuilding Word document...")
        with open(article_path, "r", encoding="utf-8") as f:
            article_md = f.read()

        doc = build_document(article_md, chart_paths)

        output_path = os.path.join(_this_dir, "UPRO_Timing_Strategies_Article.docx")
        try:
            with open(output_path, "ab") as _:
                pass
        except PermissionError:
            output_path = os.path.join(_this_dir, "UPRO_Timing_Strategies_Article_v2.docx")
            print(f"  Primary file locked, saving to: {os.path.basename(output_path)}")
        doc.save(output_path)
        print(f"\nDocument saved to: {output_path}")
    else:
        print(f"\nSkipping .docx build (article not found: {article_path})")

    print("=" * 60)


if __name__ == "__main__":
    main()
